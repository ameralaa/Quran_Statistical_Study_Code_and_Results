# advanced_analysis_and_report.py
import math
import json
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from collections import Counter
from scipy import stats
from sklearn.linear_model import LinearRegression
import docx
from docx.shared import Inches

ROOT = Path('.')
RESULTS = ROOT / 'final_results' 
OUTPUT = ROOT / RESULTS / 'step2_advanced_analysis'
OUTPUT.mkdir(exist_ok=True)

# ---------- load summary ----------
with open(RESULTS / 'step1_preprocess' / 'summary.json', 'r', encoding='utf-8') as f:
    summary = json.load(f)

# ---------- load word freq ----------
word_df = pd.read_csv(RESULTS / 'step1_preprocess' / 'word_freq.csv', encoding='utf-8')
word_df = word_df.sort_values('count', ascending=False).reset_index(drop=True)
word_df['rank'] = word_df.index + 1

# ---------- top 30 words ----------
top30 = word_df.head(30)
top30.to_csv(OUTPUT / 'top30_words.csv', index=False, encoding='utf-8')

# ---------- Zipf plot ----------
freqs = word_df['count'].values
ranks_full = word_df['rank'].values

plt.figure(figsize=(8,6))
plt.loglog(ranks_full, freqs, marker='.', linestyle='none', color='blue', alpha=0.6)
plt.xlabel('Rank (log)')
plt.ylabel('Frequency (log)')
plt.title('Zipf Law Distribution (Full Corpus)')
plt.grid(True, which="both", ls="-", alpha=0.2)
plt.tight_layout()
plt.savefig(OUTPUT / 'zipf_full.png')
plt.close()

# ---------- Fit power-law (Zipf) on middle region ----------
lo, hi = 20, min(2000, len(word_df))
x_zipf = np.log(word_df['rank'].iloc[lo:hi].values).reshape(-1,1)
y_zipf = np.log(word_df['count'].iloc[lo:hi].values)
model_zipf = LinearRegression().fit(x_zipf, y_zipf)
slope = model_zipf.coef_[0]
r2 = model_zipf.score(x_zipf, y_zipf)
fit_summary = {'slope': float(slope), 'intercept': float(model_zipf.intercept_), 'r2': float(r2), 'fit_reg_range': [int(lo+1), int(hi)]}

# Residuals analysis
pred = model_zipf.predict(x_zipf)
residuals = y_zipf - pred
ks_stat, ks_p = stats.kstest(residuals, 'norm', args=(residuals.mean(), residuals.std()))

# ---------- Shannon entropy ----------
counts_words = word_df['count'].values
p_words = counts_words / counts_words.sum()
entropy_words = -np.sum(p_words * np.log2(p_words))

char_df = pd.read_csv(RESULTS / 'step1_preprocess' / 'char_freq_without_diacritics.csv', encoding='utf-8')
counts_chars = char_df['count'].values
p_chars = counts_chars / counts_chars.sum()
entropy_chars = -np.sum(p_chars * np.log2(p_chars))

# ---------- Autocorrelation ----------
verse_df = pd.read_csv(RESULTS / 'step1_preprocess' / 'verse_metrics.csv', encoding='utf-8')
series = verse_df['n_words'].values - verse_df['n_words'].mean()
def autocorr(x, lag):
    return np.corrcoef(x[:-lag], x[lag:])[0,1]
autocorrs = {lag: autocorr(series, lag) for lag in range(1, 31)}

# ---------- Hurst Exponent & Visualization ----------
def calculate_and_plot_hurst(ts, output_path):
    N = len(ts)
    max_k = int(np.floor(np.log2(N)))
    # Create lags (sizes) on a log scale
    sizes = np.unique(np.floor(np.logspace(np.log10(10), np.log10(N/2), 15)).astype(int))
    rsvals = []
    
    for s in sizes:
        # Split series into chunks of size s
        num_chunks = N // s
        rescaled_ranges = []
        for i in range(num_chunks):
            chunk = ts[i*s : (i+1)*s]
            mean_adj = chunk - np.mean(chunk)
            cum_sum = np.cumsum(mean_adj)
            R = np.max(cum_sum) - np.min(cum_sum)
            S = np.std(chunk, ddof=1)
            if S > 0:
                rescaled_ranges.append(R / S)
        if rescaled_ranges:
            rsvals.append(np.mean(rescaled_ranges))
    
    # Fit log-log
    log_sizes = np.log(sizes[:len(rsvals)]).reshape(-1,1)
    log_rs = np.log(rsvals)
    hurst_model = LinearRegression().fit(log_sizes, log_rs)
    H = hurst_model.coef_[0]

    # Plotting
    plt.figure(figsize=(8,5))
    plt.loglog(sizes[:len(rsvals)], rsvals, 'bo', label='Measured R/S')
    plt.loglog(sizes[:len(rsvals)], np.exp(hurst_model.predict(log_sizes)), 'r-', label=f'Hurst Fit (H={H:.3f})')
    plt.xlabel('Lag (n)')
    plt.ylabel('Rescaled Range (R/S)')
    plt.title('Hurst Exponent Analysis (Rescaled Range Analysis)')
    plt.legend()
    plt.grid(True, which="both", ls="-", alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_path / 'hurst_analysis.png')
    plt.close()
    return float(H)

hurst_val = calculate_and_plot_hurst(verse_df['n_words'].values, OUTPUT)

# ---------- Save Summary JSON ----------
adv_summary = {
    'zipf_fit': fit_summary,
    'zipf_fit_ks_residuals': {'ks_stat': float(ks_stat), 'ks_p': float(ks_p)},
    'entropy_words': float(entropy_words),
    'entropy_chars': float(entropy_chars),
    'autocorr_first_10': {str(k): float(v) for k,v in list(autocorrs.items())[:10]},
    'hurst': hurst_val
}
with open(OUTPUT / 'advanced_summary.json', 'w', encoding='utf-8') as f:
    json.dump(adv_summary, f, ensure_ascii=False, indent=2)

# ---------- Assemble English DOCX Report ----------
doc = docx.Document()
doc.add_heading('Advanced Statistical Analysis of the Quranic Text', level=1)
doc.add_paragraph(f"Source: Preprocessed data from results/ folder. Total Verses: {summary['n_verses']}, Total Tokens: {summary['n_tokens']}.")

doc.add_heading('1. Zipf Law Analysis', level=2)
doc.add_paragraph(f"A linear model was fitted to log-ranks and log-frequencies in the range {fit_summary['fit_reg_range']}.")
doc.add_paragraph(f"Slope (alpha): {fit_summary['slope']:.4f}, R-squared: {fit_summary['r2']:.4f}")
doc.add_picture(str(OUTPUT / 'zipf_full.png'), width=Inches(5.5))

doc.add_heading('2. Information Theory Metrics (Entropy)', level=2)
doc.add_paragraph(f"Shannon Entropy (Words): {entropy_words:.4f} bits")
doc.add_paragraph(f"Shannon Entropy (Characters): {entropy_chars:.4f} bits")

doc.add_heading('3. Long-Range Dependency & Hurst Exponent', level=2)
doc.add_paragraph(f"Estimated Hurst Exponent (H): {hurst_val:.4f}")
doc.add_paragraph("Interpretation: A value H > 0.5 indicates strong persistence and long-range structural memory.")
doc.add_picture(str(OUTPUT / 'hurst_analysis.png'), width=Inches(5.5))

doc.add_heading('4. Autocorrelation Analysis', level=2)
doc.add_paragraph("Autocorrelation of verse lengths (lags 1-10):")
for k in range(1,11):
    doc.add_paragraph(f"Lag {k}: {autocorrs[k]:.4f}")

doc.add_heading('5. Statistical Residuals & Anomaly Notes', level=2)
doc.add_paragraph("Full numerical data is available in 'advanced_summary.json'. The Kolmogorov-Smirnov test on Zipf residuals yielded:")
doc.add_paragraph(f"KS-Statistic: {ks_stat:.4f}, p-value: {ks_p:.4f}")

doc.save(OUTPUT / 'advanced_report.docx')

print("Done. English report and Hurst plot generated in:", OUTPUT)
